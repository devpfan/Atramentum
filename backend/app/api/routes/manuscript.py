from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional
from app.db.database import get_db
from app.models.manuscript import Book, Act, Chapter, Scene
from app.models.user import User
from app.schemas.manuscript import ManuscriptTree, Chapter as ChapterSchema, Scene as SceneSchema, ChapterCreate, SceneCreate, SceneUpdate
from pydantic import BaseModel
from datetime import datetime
from app.api.deps import get_current_user
from app.services.importer import import_document
from app.services.vector_store import sync_scene_embeddings

router = APIRouter()

class BookBase(BaseModel):
    title: str
    synopsis: Optional[str] = None
    target_word_count: Optional[int] = 50000
    series_id: Optional[int] = None

class BookSchema(BookBase):
    id: int
    user_id: int
    created_at: datetime
    series_id: Optional[int] = None
    
    class Config:
        from_attributes = True

@router.get("/books", response_model=List[BookSchema])
def get_books(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    books = db.query(Book).filter(Book.user_id == current_user.id).order_by(Book.created_at.desc()).all()
    if not books:
        # Create default book if none exists
        book = Book(title="Proyecto Sin Título", user_id=current_user.id)
        db.add(book)
        db.commit()
        db.refresh(book)
        
        act = Act(title="Acto 1", book_id=book.id, user_id=current_user.id)
        db.add(act)
        db.commit()
        db.refresh(act)
        
        chapter = Chapter(title="Capítulo 1", act_id=act.id, user_id=current_user.id)
        db.add(chapter)
        db.commit()
        db.refresh(chapter)
        
        scene = Scene(
            title="Escena 1", 
            chapter_id=chapter.id, 
            user_id=current_user.id,
            content="<p>Érase una vez...</p>"
        )
        db.add(scene)
        db.commit()
        
        books = [book]
    return books

@router.post("/books", response_model=BookSchema)
def create_book(book_in: BookBase, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    book = Book(
        title=book_in.title, 
        synopsis=book_in.synopsis, 
        target_word_count=book_in.target_word_count,
        series_id=book_in.series_id,
        user_id=current_user.id
    )
    db.add(book)
    db.commit()
    db.refresh(book)
    
    act = Act(title="Acto 1", book_id=book.id, user_id=current_user.id)
    db.add(act)
    db.commit()
    db.refresh(act)
    
    chapter = Chapter(title="Capítulo 1", act_id=act.id, user_id=current_user.id)
    db.add(chapter)
    db.commit()
    db.refresh(chapter)
    
    scene = Scene(
        title="Escena 1", 
        chapter_id=chapter.id, 
        user_id=current_user.id,
        content="<p>Érase una vez...</p>"
    )
    db.add(scene)
    db.commit()
    
    return book

@router.post("/import", response_model=ManuscriptTree)
async def import_book(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    try:
        content = await file.read()
        parsed_chapters = import_document(content, file.filename)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    title = file.filename.rsplit('.', 1)[0]
    
    book = Book(
        title=title,
        user_id=current_user.id
    )
    db.add(book)
    db.commit()
    db.refresh(book)
    
    act = Act(title="Acto 1", book_id=book.id, user_id=current_user.id)
    db.add(act)
    db.commit()
    db.refresh(act)
    
    for i, ch_data in enumerate(parsed_chapters, start=1):
        chapter = Chapter(
            title=ch_data["title"],
            order=i,
            act_id=act.id,
            user_id=current_user.id
        )
        db.add(chapter)
        db.flush()  # To get chapter.id without committing yet
        
        scene = Scene(
            title="Escena 1",
            order=1,
            content=ch_data["html_content"],
            chapter_id=chapter.id,
            user_id=current_user.id
        )
        db.add(scene)
        
    db.commit()
    
    # Reload tree
    acts = db.query(Act).filter(Act.book_id == book.id).all()
    act_ids = [a.id for a in acts]
    chapters = db.query(Chapter).filter(Chapter.act_id.in_(act_ids)).order_by(Chapter.order, Chapter.id).all()
    
    return ManuscriptTree(
        book_id=book.id,
        title=book.title,
        chapters=chapters
    )

class BookUpdate(BaseModel):
    title: Optional[str] = None
    synopsis: Optional[str] = None
    target_word_count: Optional[int] = None
    series_id: Optional[int] = None

@router.put("/books/{book_id}", response_model=BookSchema)
def update_book(book_id: int, book_update: BookUpdate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    from sqlalchemy.orm.attributes import flag_modified
    book = db.query(Book).filter(Book.id == book_id, Book.user_id == current_user.id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Libro no encontrado")
        
    if book_update.title and book_update.title != book.title:
        # Añadir el título antiguo al historial si no es nulo
        if book.previous_titles is None:
            book.previous_titles = []
        book.previous_titles.append(book.title)
        flag_modified(book, "previous_titles")
        book.title = book_update.title
        
    if book_update.synopsis is not None:
        book.synopsis = book_update.synopsis
        
    if book_update.target_word_count is not None:
        book.target_word_count = book_update.target_word_count
        
    if book_update.series_id is not None:
        book.series_id = book_update.series_id
        
    db.commit()
    db.refresh(book)
    return book

@router.delete("/books/{book_id}")
def delete_book(book_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    book = db.query(Book).filter(Book.id == book_id, Book.user_id == current_user.id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Libro no encontrado")
    
    db.delete(book)
    db.commit()
    return {"message": "Libro eliminado con éxito"}

@router.get("/tree", response_model=ManuscriptTree)
def get_manuscript_tree(book_id: Optional[int] = None, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """
    Devuelve la estructura completa de Capítulos y Escenas.
    Si el usuario no tiene ningún libro, inicializa uno por defecto.
    """
    if book_id:
        book = db.query(Book).filter(Book.id == book_id, Book.user_id == current_user.id).first()
        if not book:
            raise HTTPException(status_code=404, detail="Libro no encontrado")
    else:
        book = db.query(Book).filter(Book.user_id == current_user.id).order_by(Book.created_at.desc()).first()
    
    if not book:
        # Inicialización por defecto
        book = Book(title="Proyecto Sin Título", user_id=current_user.id)
        db.add(book)
        db.commit()
        db.refresh(book)
        
        act = Act(title="Acto 1", book_id=book.id, user_id=current_user.id)
        db.add(act)
        db.commit()
        db.refresh(act)
        
        chapter = Chapter(title="Capítulo 1", act_id=act.id, user_id=current_user.id)
        db.add(chapter)
        db.commit()
        db.refresh(chapter)
        
        scene = Scene(
            title="Escena 1", 
            chapter_id=chapter.id, 
            user_id=current_user.id,
            content="<p>Érase una vez...</p>"
        )
        db.add(scene)
        db.commit()
        db.refresh(scene)
    
    # Para el UI actual, aplanamos los Capítulos de todos los Actos de este Libro.
    acts = db.query(Act).filter(Act.book_id == book.id).all()
    act_ids = [a.id for a in acts]
    
    chapters = db.query(Chapter).filter(Chapter.act_id.in_(act_ids)).order_by(Chapter.order, Chapter.id).all()
    
    return ManuscriptTree(
        book_id=book.id,
        title=book.title,
        chapters=chapters
    )

@router.post("/chapters", response_model=ChapterSchema)
def create_chapter(chapter: ChapterCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    act_id = chapter.act_id
    if not act_id:
        if chapter.book_id:
            act = db.query(Act).filter(Act.book_id == chapter.book_id, Act.user_id == current_user.id).first()
            if not act:
                act = Act(title="Acto 1", book_id=chapter.book_id, user_id=current_user.id)
                db.add(act)
                db.commit()
                db.refresh(act)
            act_id = act.id
        else:
            act = db.query(Act).filter(Act.user_id == current_user.id).first()
            if act:
                act_id = act.id

    if not act_id:
        raise HTTPException(status_code=400, detail="Acto o Libro requerido para crear capítulo")

    db_chapter = Chapter(
        title=chapter.title, 
        order=chapter.order or 1, 
        act_id=act_id, 
        user_id=current_user.id
    )
    db.add(db_chapter)
    db.commit()
    db.refresh(db_chapter)
    return db_chapter

class ChapterUpdate(BaseModel):
    title: Optional[str] = None
    order: Optional[int] = None

@router.put("/chapters/{chapter_id}", response_model=ChapterSchema)
def update_chapter(chapter_id: int, chapter_update: ChapterUpdate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    db_chapter = db.query(Chapter).filter(Chapter.id == chapter_id, Chapter.user_id == current_user.id).first()
    if not db_chapter:
        raise HTTPException(status_code=404, detail="Capítulo no encontrado")
    
    update_data = chapter_update.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_chapter, key, value)
        
    db.commit()
    db.refresh(db_chapter)
    return db_chapter

@router.delete("/chapters/{chapter_id}")
def delete_chapter(chapter_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    chapter = db.query(Chapter).filter(Chapter.id == chapter_id, Chapter.user_id == current_user.id).first()
    if not chapter:
        raise HTTPException(status_code=404, detail="Capítulo no encontrado")
    db.delete(chapter)
    db.commit()
    return {"message": "Capítulo eliminado exitosamente"}

@router.delete("/scenes/{scene_id}")
def delete_scene(scene_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    scene = db.query(Scene).filter(Scene.id == scene_id, Scene.user_id == current_user.id).first()
    if not scene:
        raise HTTPException(status_code=404, detail="Escena no encontrada")
    db.delete(scene)
    db.commit()
    return {"message": "Escena eliminada exitosamente"}

@router.put("/scenes/{scene_id}", response_model=SceneSchema)
def update_scene(
    scene_id: int, 
    scene_in: SceneUpdate, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    scene = db.query(Scene).filter(Scene.id == scene_id, Scene.user_id == current_user.id).first()
    if not scene:
        raise HTTPException(status_code=404, detail="Scene not found")
        
    # Validar que si intentamos cambiar de chapter, este nos pertenezca
    if scene_in.chapter_id is not None:
        new_chap = db.query(Chapter).filter(Chapter.id == scene_in.chapter_id, Chapter.user_id == current_user.id).first()
        if not new_chap:
            raise HTTPException(status_code=404, detail="Target Chapter not found")
            
    update_data = scene_in.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(scene, field, value)
        
    db.commit()
    db.refresh(scene)
    
    # Si se actualizó el contenido, regeneramos los embeddings en background
    if "content" in update_data:
        background_tasks.add_task(sync_scene_embeddings, scene.id, current_user.id, db)
        
    return scene

@router.post("/scenes", response_model=SceneSchema)
def create_scene(scene: SceneCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    db_scene = Scene(**scene.model_dump(), user_id=current_user.id)
    db.add(db_scene)
    db.commit()
    db.refresh(db_scene)
    return db_scene

class ReorderItem(BaseModel):
    id: int
    type: str # "act", "chapter", "scene"
    order: int
    parent_id: Optional[int] = None # act_id for chapter, chapter_id for scene

class BulkReorderRequest(BaseModel):
    items: List[ReorderItem]

@router.put("/books/{book_id}/reorder")
def bulk_reorder(book_id: int, req: BulkReorderRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    # Very simple bulk reorder
    for item in req.items:
        if item.type == "chapter":
            db_item = db.query(Chapter).filter(Chapter.id == item.id, Chapter.user_id == current_user.id).first()
            if db_item:
                db_item.order = item.order
                if item.parent_id is not None:
                    db_item.act_id = item.parent_id
        elif item.type == "scene":
            db_item = db.query(Scene).filter(Scene.id == item.id, Scene.user_id == current_user.id).first()
            if db_item:
                db_item.order = item.order
                if item.parent_id is not None:
                    db_item.chapter_id = item.parent_id
    db.commit()
    return {"message": "Reordered successfully"}

from fastapi.responses import StreamingResponse
from io import BytesIO, StringIO
import io

@router.get("/books/{book_id}/export")
def export_book(book_id: int, format: str = "md", db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    book = db.query(Book).filter(Book.id == book_id, Book.user_id == current_user.id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Libro no encontrado")

    acts = db.query(Act).filter(Act.book_id == book.id).all()
    act_ids = [a.id for a in acts]
    chapters = db.query(Chapter).filter(Chapter.act_id.in_(act_ids)).order_by(Chapter.order, Chapter.id).all()
    
    # Pre-fetch all scenes
    from sqlalchemy.orm import joinedload
    chapters = db.query(Chapter).options(joinedload(Chapter.scenes)).filter(Chapter.act_id.in_(act_ids)).order_by(Chapter.order, Chapter.id).all()

    if format == "md":
        from markdownify import markdownify as md
        content = f"# {book.title}\n\n"
        for ch in chapters:
            content += f"## {ch.title}\n\n"
            for sc in sorted(ch.scenes, key=lambda x: (x.order, x.id)):
                content += f"### {sc.title}\n\n"
                if sc.content:
                    content += md(sc.content) + "\n\n"
        
        file_stream = StringIO(content)
        return StreamingResponse(
            iter([file_stream.getvalue()]), 
            media_type="text/markdown", 
            headers={"Content-Disposition": f"attachment; filename=\"{book.title}.md\""}
        )

    elif format == "docx":
        from docx import Document
        from bs4 import BeautifulSoup
        doc = Document()
        doc.add_heading(book.title, 0)
        
        for ch in chapters:
            doc.add_heading(ch.title, 1)
            for sc in sorted(ch.scenes, key=lambda x: (x.order, x.id)):
                doc.add_heading(sc.title, 2)
                if sc.content:
                    soup = BeautifulSoup(sc.content, 'html.parser')
                    for p in soup.find_all('p'):
                        doc.add_paragraph(p.get_text())
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": f"attachment; filename=\"{book.title}.docx\""}
        )

    elif format == "pdf":
        from fpdf import FPDF
        from bs4 import BeautifulSoup
        class PDF(FPDF):
            def header(self):
                self.set_font('helvetica', 'I', 8)
                self.cell(0, 10, f'{book.title}', 0, 1, 'C')
            def footer(self):
                self.set_y(-15)
                self.set_font('helvetica', 'I', 8)
                self.cell(0, 10, f'Página {self.page_no()}', 0, 0, 'C')

        pdf = PDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("helvetica", "B", 24)
        pdf.cell(0, 20, book.title, 0, 1, 'C')
        pdf.ln(20)
        
        for ch in chapters:
            pdf.set_font("helvetica", "B", 18)
            pdf.cell(0, 15, ch.title, 0, 1, 'L')
            for sc in sorted(ch.scenes, key=lambda x: (x.order, x.id)):
                pdf.set_font("helvetica", "B", 14)
                pdf.cell(0, 10, sc.title, 0, 1, 'L')
                pdf.set_font("helvetica", "", 12)
                if sc.content:
                    soup = BeautifulSoup(sc.content, 'html.parser')
                    for p in soup.find_all('p'):
                        text = p.get_text().encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 8, text)
                        pdf.ln(2)
                pdf.ln(5)
                
        file_stream = BytesIO(pdf.output())
        file_stream.seek(0)
        return StreamingResponse(
            file_stream, 
            media_type="application/pdf", 
            headers={"Content-Disposition": f"attachment; filename=\"{book.title}.pdf\""}
        )

    elif format == "epub":
        from ebooklib import epub
        from bs4 import BeautifulSoup
        
        book_epub = epub.EpubBook()
        book_epub.set_identifier(f"atramentum-{book.id}")
        book_epub.set_title(book.title)
        book_epub.set_language('es')
        book_epub.add_author("Autor") # Could be current_user.name
        
        chapters_epub = []
        for i, ch in enumerate(chapters):
            c = epub.EpubHtml(title=ch.title, file_name=f'chap_{i}.xhtml', lang='es')
            content_html = f"<h1>{ch.title}</h1>"
            for sc in sorted(ch.scenes, key=lambda x: (x.order, x.id)):
                content_html += f"<h2>{sc.title}</h2>"
                if sc.content:
                    soup = BeautifulSoup(sc.content, 'html.parser')
                    for p in soup.find_all('p'):
                        content_html += f"<p>{p.get_text()}</p>"
            
            c.content = content_html
            book_epub.add_item(c)
            chapters_epub.append(c)
            
        book_epub.toc = tuple(chapters_epub)
        book_epub.add_item(epub.EpubNcx())
        book_epub.add_item(epub.EpubNav())
        
        # Add default style
        style = 'BODY {color: white;}'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book_epub.add_item(nav_css)
        book_epub.spine = ['nav'] + chapters_epub
        
        file_stream = BytesIO()
        epub.write_epub(file_stream, book_epub)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream, 
            media_type="application/epub+zip", 
            headers={"Content-Disposition": f"attachment; filename=\"{book.title}.epub\""}
        )

    else:
        raise HTTPException(status_code=400, detail="Formato no soportado")

